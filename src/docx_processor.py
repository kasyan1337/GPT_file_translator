# src/docx_processor.py

from docx import Document
from src.document_processor import DocumentProcessor
from tqdm import tqdm
from src.utils import chunk_runs

class DocxProcessor(DocumentProcessor):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.total_usage = {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
        }

    def process(self):
        document = Document(self.file_path)

        # Collect all runs with text in the document, including tables, headers, footers, etc.
        runs = []
        # Process paragraphs
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    runs.append(run)
        # Process tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                runs.append(run)
        # Process headers and footers
        for section in document.sections:
            header = section.header
            footer = section.footer
            for paragraph in header.paragraphs + footer.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        runs.append(run)

        # Group runs into chunks
        run_chunks = chunk_runs(
            runs, max_tokens=7500, model=self.openai_api.model
        )

        for chunk in tqdm(run_chunks, desc="Translating DOCX"):
            original_text = "\n".join([run.text for run in chunk])
            translated_text = self.openai_api.translate_text(self.prompt, original_text)
            if translated_text:
                translated_runs = translated_text.split("\n")
                for run, t in zip(chunk, translated_runs):
                    run.text = t
                usage = self.openai_api.last_usage
                self.total_usage["prompt_tokens"] += usage["prompt_tokens"]
                self.total_usage["completion_tokens"] += usage["completion_tokens"]
                self.total_usage["total_tokens"] += usage["total_tokens"]
            else:
                print("Translation failed for a chunk.")

        document.save(self.output_path)

        print(f"Token usage for {self.file_path}: {self.total_usage}")
        estimated_cost = self.openai_api.calculate_cost(self.total_usage)
        print(f"Estimated cost for {self.file_path}: ${estimated_cost:.4f}")